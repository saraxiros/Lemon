import io
import logging
import os
import tempfile
import traceback
from concurrent.futures import ThreadPoolExecutor, as_completed

import google.generativeai as genai
from dotenv import load_dotenv
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.requests import Request
from fastapi.responses import FileResponse, JSONResponse
from PyPDF2 import PdfReader, PdfWriter

# =====================================================================
# Logging — technical details are logged here, never sent to the client.
# =====================================================================

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger("doc-converter")

# =====================================================================
# Configuration
# =====================================================================

FILES_FOLDER = "files"      # kept for backward compatibility / manual batch use
CURATED_FOLDER = "curated"  # generated .docx files are also archived here

MIN_ADDITIONAL_ROWS = 0
MAX_ADDITIONAL_ROWS = 200

DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# Comma-separated list of allowed frontend origins, configurable via .env.
# Example .env line: ALLOWED_ORIGINS=http://localhost:5500,http://127.0.0.1:5500
DEFAULT_ALLOWED_ORIGINS = "http://localhost:5500,http://127.0.0.1:5500"


def get_gemini_api_key():
    """Loads the Gemini API key from the .env file."""
    load_dotenv()
    return os.getenv("GEMINI_API_KEY")


def get_allowed_origins():
    load_dotenv()
    raw = os.getenv("ALLOWED_ORIGINS", DEFAULT_ALLOWED_ORIGINS)
    return [origin.strip() for origin in raw.split(",") if origin.strip()]


# =====================================================================
# Core OCR / document-generation pipeline
# (unchanged from the original script, aside from removing input())
# =====================================================================

def extract_text_from_pdf_page(page_index, pdf_page_bytes, model):
    """
    Uses the Gemini API to extract text from a single PDF page's bytes.
    Returns tuple: (page_index, text)
    """
    uploaded_file = None
    tmp_path = None

    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(pdf_page_bytes)
            tmp_path = tmp.name

        uploaded_file = genai.upload_file(
            path=tmp_path,
            display_name=f"page_{page_index + 1}.pdf",
            mime_type="application/pdf",
        )

        prompt = """
        Extract only the relevant multiple-choice content from this PDF page. Specifically, include:
        - The module name, if it is mentioned on the page, written alone (one word) and in capital letters. If not present, do not invent one.
        - The questions and their answer options only.
        - Exclude all weird marks or symbols next to MCQs.
        - Exclude instructions, titles, headers, footers, and page numbers.
        - Preserve the original question numbers but format them using two digits (01., 02., 03., etc.).
        - Convert all answer option labels to: a) b) c) d) regardless of their original formatting.
        """

        response = model.generate_content([prompt, uploaded_file])

        if hasattr(response, "text") and response.text:
            text = response.text
        elif hasattr(response, "candidates") and response.candidates:
            text = response.candidates[0].content.parts[0].text
        else:
            text = "[No text extracted]"

        return (page_index, text)

    except Exception as e:
        return (page_index, f"[Error: {e}]")

    finally:
        if uploaded_file:
            try:
                genai.delete_file(uploaded_file.name)
            except Exception:
                pass

        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)


def build_docx_from_pdf(pdf_path, output_path, model, additional_rows, source_filename):
    """
    Runs the full PDF -> Gemini OCR -> .docx pipeline for a single PDF.

    This is the reusable core previously embedded in process_single_pdf().
    It no longer calls input() — the row count is passed in directly so it
    can come from an HTTP request, a CLI prompt, or anywhere else.

    Raises on failure instead of returning False, so callers (the API route,
    or a CLI wrapper) can decide how to report the error.
    """
    logger.info("Processing: %s", source_filename)

    document = Document()

    # Narrow margins
    for section in document.sections:
        section.left_margin = Pt(36)
        section.right_margin = Pt(36)
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)

    # Default font
    style = document.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.space_before = Pt(0)
    paragraph_format.line_spacing = 1

    document.add_heading(f"Extracted Text from {source_filename}", 0)

    with open(pdf_path, "rb") as f:
        reader = PdfReader(f)
        num_pages = len(reader.pages)
        logger.info("  Total pages: %d", num_pages)

        pages_bytes = []
        for i in range(num_pages):
            writer = PdfWriter()
            writer.add_page(reader.pages[i])
            with io.BytesIO() as bytes_stream:
                writer.write(bytes_stream)
                pages_bytes.append((i, bytes_stream.getvalue()))

        # Process pages in parallel
        results = {}
        with ThreadPoolExecutor(max_workers=5) as executor:
            futures = {
                executor.submit(extract_text_from_pdf_page, i, data, model): i
                for i, data in pages_bytes
            }

            for future in as_completed(futures):
                i, text = future.result()
                results[i] = text
                logger.info("    Page %d/%d done", i + 1, num_pages)

        # Write results in correct order
        for i in range(num_pages):
            document.add_heading(f"--- Page {i + 1} ---", level=1)

            paragraph = document.add_paragraph()

            for line in results[i].split("\n"):
                if line.strip().startswith(tuple(f"{n:02d}." for n in range(1, 101))):
                    run = paragraph.add_run(line + "\n")
                    run.bold = True
                else:
                    paragraph.add_run(line + "\n")

            document.add_page_break()

        # Corrigé Type
        title = document.add_paragraph()
        run = title.add_run("Corrigé Type")
        run.bold = True
        run.italic = True
        run.underline = True
        title.alignment = 1

        document.add_paragraph()

        table = document.add_table(rows=additional_rows, cols=4)
        table.style = "Table Grid"

        for i in range(additional_rows):
            left = f"Q{i + 1:02d}"
            right = f"Q{additional_rows + i + 1:02d}"

            table.cell(i, 0).text = left
            table.cell(i, 2).text = right

        document.save(output_path)
        logger.info("  Saved to: %s", output_path)


# =====================================================================
# Gemini model setup
# =====================================================================

model = None


def init_gemini_model():
    global model
    api_key = get_gemini_api_key()
    if not api_key:
        raise RuntimeError(
            "GEMINI_API_KEY not found. Add it to your .env file before starting the server."
        )

    genai.configure(api_key=api_key)

    try:
        model = genai.GenerativeModel("gemini-2.5-flash")
    except Exception:
        logger.warning("Falling back to gemini-pro")
        model = genai.GenerativeModel("gemini-pro")


# =====================================================================
# FastAPI app
# =====================================================================

app = FastAPI(title="Doc Converter API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=get_allowed_origins(),
    allow_credentials=True,
    allow_methods=["POST", "GET", "OPTIONS"],
    allow_headers=["*"],
)


@app.on_event("startup")
def on_startup():
    os.makedirs(FILES_FOLDER, exist_ok=True)
    os.makedirs(CURATED_FOLDER, exist_ok=True)
    init_gemini_model()
    logger.info("Doc Converter API ready.")


@app.exception_handler(Exception)
async def unhandled_exception_handler(request: Request, exc: Exception):
    # Catch-all: never leak stack traces / internals to the client.
    logger.error("Unhandled error on %s: %s\n%s", request.url.path, exc, traceback.format_exc())
    return JSONResponse(status_code=500, content={"detail": "Something went wrong on our end."})


@app.get("/")
def root():
    return {"status": "ok"}


@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/convert")
def convert(pdf: UploadFile = File(...), additional_rows: int = Form(...)):
    """
    Accepts a PDF + additional_rows from the frontend, runs the existing
    Gemini OCR / docx pipeline, and returns the generated .docx directly.
    """

    # --- validate the upload ---
    filename = pdf.filename or "document.pdf"
    looks_like_pdf_type = pdf.content_type == "application/pdf"
    looks_like_pdf_name = filename.lower().endswith(".pdf")

    if not looks_like_pdf_type and not looks_like_pdf_name:
        raise HTTPException(status_code=400, detail="Please upload a PDF file.")

    # --- validate additional_rows ---
    if additional_rows < MIN_ADDITIONAL_ROWS or additional_rows > MAX_ADDITIONAL_ROWS:
        raise HTTPException(
            status_code=400,
            detail=f"additional_rows must be between {MIN_ADDITIONAL_ROWS} and {MAX_ADDITIONAL_ROWS}.",
        )

    if model is None:
        raise HTTPException(status_code=503, detail="The conversion service is not ready yet.")

    # --- save the upload to a temporary file (existing pipeline needs a path) ---
    tmp_pdf_path = None
    try:
        content = pdf.file.read()
        if not content:
            raise HTTPException(status_code=400, detail="The uploaded file is empty.")

        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(content)
            tmp_pdf_path = tmp.name
    except HTTPException:
        raise
    except Exception:
        logger.exception("Failed to save uploaded PDF")
        raise HTTPException(status_code=400, detail="Unable to read the uploaded PDF.")
    finally:
        pdf.file.close()

    base_name = os.path.splitext(os.path.basename(filename))[0] or "document"
    output_filename = f"curated_{base_name}.docx"
    output_path = os.path.join(CURATED_FOLDER, output_filename)

    try:
        build_docx_from_pdf(tmp_pdf_path, output_path, model, additional_rows, filename)
    except Exception:
        logger.exception("Conversion pipeline failed for %s", filename)
        raise HTTPException(status_code=500, detail="Unable to convert this document. Please try again.")
    finally:
        if tmp_pdf_path and os.path.exists(tmp_pdf_path):
            os.remove(tmp_pdf_path)

    if not os.path.exists(output_path):
        raise HTTPException(status_code=500, detail="Conversion did not produce a document.")

    return FileResponse(
        path=output_path,
        media_type=DOCX_MIME_TYPE,
        filename=output_filename,
    )


# =====================================================================
# Optional CLI batch mode — preserves the original folder-based workflow.
# Run with `python main.py` (not through uvicorn) to batch-process
# everything in files/ the same way the original script did.
# =====================================================================

def process_single_pdf_cli(pdf_path, output_path, model_instance):
    filename = os.path.basename(pdf_path)
    rows = int(input(f"Enter number of rows you want for {filename}: "))
    build_docx_from_pdf(pdf_path, output_path, model_instance, rows, filename)


def run_batch_cli():
    init_gemini_model()

    if not os.path.exists(FILES_FOLDER):
        print(f"Error: '{FILES_FOLDER}' folder not found.")
        return

    os.makedirs(CURATED_FOLDER, exist_ok=True)

    pdf_files = [f for f in os.listdir(FILES_FOLDER) if f.lower().endswith(".pdf")]

    if not pdf_files:
        print(f"No PDF files found in '{FILES_FOLDER}'.")
        return

    print(f"Found {len(pdf_files)} PDF file(s) to process.\n")

    successful = 0
    failed = 0

    for pdf_file in pdf_files:
        pdf_path = os.path.join(FILES_FOLDER, pdf_file)
        base_name = os.path.splitext(pdf_file)[0]
        output_filename = f"curated_{base_name}.docx"
        output_path = os.path.join(CURATED_FOLDER, output_filename)

        try:
            process_single_pdf_cli(pdf_path, output_path, model)
            successful += 1
        except Exception as e:
            print(f"  Error processing {pdf_file}: {e}")
            failed += 1

    print("\n" + "=" * 50)
    print(f"Successfully processed: {successful} file(s)")
    if failed:
        print(f"Failed: {failed} file(s)")
    print(f"Output location: {CURATED_FOLDER}/")
    print("=" * 50)


if __name__ == "__main__":
    run_batch_cli()
